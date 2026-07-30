import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Page Margins: Top 2.0cm, Bottom 2.0cm, Left 3.0cm, Right 2.0cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)     # ~2.0 cm
        section.bottom_margin = Inches(0.79)  # ~2.0 cm
        section.left_margin = Inches(1.18)    # ~3.0 cm
        section.right_margin = Inches(0.79)   # ~2.0 cm

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(13)
    normal_font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(4)

    # Helper Functions
    def add_p(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=13, space_before=0, space_after=4, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.25
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark Navy
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153) # Navy Accent
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
        return p

    def set_cell_background(cell, fill_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        tcPr.append(shd)

    # -------------------------------------------------------------
    # TRANG BÌA (COVER PAGE - MẪU M1 CHUẨN KHOA CSE - ĐHBK)
    # -------------------------------------------------------------
    add_p("TRƯỜNG ĐẠI HỌC BÁCH KHOA - ĐHQG TP. HỒ CHÍ MINH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=2)
    add_p("KHOA KHOA HỌC VÀ KỸ THUẬT MÁY TÍNH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=12)
    
    add_p("—-------------------------------------------------------------", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=16)

    # Insert BK Logo (M2_Logo_BK.png)
    logo_path = r"e:\AWS - TTNT\M2_Logo_BK.png"
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(6)
        p_logo.paragraph_format.space_after = Pt(16)
        p_logo.add_run().add_picture(logo_path, width=Inches(2.1))

    add_p("BÁO CÁO MÔN HỌC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=4, color=RGBColor(0, 51, 102))
    add_p("THỰC TẬP NGOÀI TRƯỜNG (TTNT)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, space_after=6, color=RGBColor(0, 51, 102))
    add_p("MÃ SỐ MÔN HỌC: CO3011", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=20)

    # Project Title Box
    add_p("ĐỀ TÀI THỰC TẬP:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=2)
    add_p("XÂY DỰNG HỆ THỐNG DỰ BÁO DOANH SỐ BÁN LẺ VỚI AMAZON SAGEMAKER VÀ KIẾN TRÚC MLOPS TRÊN HẠ TẦNG AWS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13.5, space_after=24, color=RGBColor(153, 0, 0))

    # Meta Info Table
    meta_table = doc.add_table(rows=8, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    col_widths = [Inches(2.5), Inches(4.0)]
    meta_data = [
        ("Học kỳ / Năm học:", "Học kỳ 3 (Hè) — Năm học 2025 - 2026"),
        ("Ngành học:", "Khoa học Máy tính"),
        ("Chương trình đào tạo:", "Chương trình Chính quy (CQ)"),
        ("Doanh nghiệp tiếp nhận:", "Amazon Web Services (AWS Viet Nam Company Limited)"),
        ("Chương trình thực tập:", "First Cloud Journey (FCJ) — AWS AI/ML Track"),
        ("Cán bộ hướng dẫn (DN):", "Nguyễn Ngọc Sáng / Nguyễn Đức Hoàng"),
        ("Cán bộ hỗ trợ (Khoa):", "TS. Nguyễn Đức Dũng / ThS. Vũ Văn Tiến"),
        ("Sinh viên thực hiện:", "Huỳnh Kim Quý — MSSV: 2312918")
    ]

    for i, (label, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        cell_lbl = row.cells[0]
        cell_lbl.width = col_widths[0]
        p0 = cell_lbl.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.bold = True
        
        cell_val = row.cells[1]
        cell_val.width = col_widths[1]
        p1 = cell_val.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(val)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        if label.startswith("Sinh viên"):
            r1.bold = True
            r1.font.color.rgb = RGBColor(0, 51, 102)

    add_p("", space_after=30)
    add_p("TP. HỒ CHÍ MINH, THÁNG 08 NĂM 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)

    doc.add_page_break()

    # -------------------------------------------------------------
    # MỤC LỤC & LỜI CAM ĐOAN
    # -------------------------------------------------------------
    add_h1("LỜI CAM ĐOAN VÀ XÁC NHẬN BẢO MẬT")
    add_p("Tôi xin cam đoan báo cáo thực tập này là công trình thực hiện thực tế của tôi cùng các thành viên trong nhóm thực tập dưới sự hướng dẫn chuyên môn của Cán bộ hướng dẫn tại Doanh nghiệp Amazon Web Services (AWS) và Cán bộ hỗ trợ của Khoa Khoa học và Kỹ thuật Máy tính - Trường Đại học Bách Khoa TP.HCM.")
    add_p("Các thông tin, dữ liệu, kết quả thực nghiệm và mã nguồn được trình bày trong báo cáo là trung thực, tuân thủ nghiêm ngặt các quy định về an toàn thông tin và bảo mật dữ liệu của Doanh nghiệp tiếp nhận thực tập.")
    
    add_p("", space_after=10)

    add_h1("MỤC LỤC BÁO CÁO")
    toc_items = [
        ("PHẦN 1: CHƯƠNG TRÌNH THỰC TẬP VÀ THÔNG TIN DOANH NGHIỆP", "Trang 3"),
        ("  1.1. Giới thiệu Doanh nghiệp tiếp nhận (Amazon Web Services)", "Trang 3"),
        ("  1.2. Chương trình thực tập đã được Khoa duyệt (Form D2/D3)", "Trang 3"),
        ("PHẦN 2: NỘI DUNG VÀ NHẬT KÝ HOẠT ĐỘNG THỰC TẬP (WORKLOG)", "Trang 4"),
        ("  2.1. Nhật ký công việc từng tuần (Tuần 1 → Tuần 8)", "Trang 4"),
        ("  2.2. Chi tiết kết quả thực hiện các giai đoạn", "Trang 6"),
        ("PHẦN 3: BÁO CÁO KỸ THUẬT VÀ KIẾN TRÚC GIẢI PHÁP", "Trang 8"),
        ("  3.1. Tổng quan Kiến trúc MLOps 4 Tầng", "Trang 8"),
        ("  3.2. Tiền xử lý dữ liệu và Feature Engineering (22 đặc trưng)", "Trang 9"),
        ("  3.3. Huấn luyện và So sánh Mô hình (XGBoost vs PyTorch LSTM)", "Trang 11"),
        ("  3.4. Quản lý Thí nghiệm với Amazon SageMaker Experiments", "Trang 13"),
        ("  3.5. Tự động hóa Pipeline với Amazon SageMaker Pipelines", "Trang 15"),
        ("  3.6. Triển khai Serverless REST API và Live Web UI Dashboard", "Trang 17"),
        ("PHẦN 4: HƯỚNG DẪN THỰC HÀNH WORKSHOP PIPELINE", "Trang 19"),
        ("PHẦN 5: TỰ ĐÁNH GIÁ VÀ PHẢN HỒI THỰC TẬP (8 TUẦN)", "Trang 23"),
        ("  5.1. Bảng Tiêu chí Đánh giá Cá nhân", "Trang 23"),
        ("  5.2. Đánh giá Tổng quan & Phản hồi Thực tập", "Trang 24"),
        ("  5.3. Đề xuất & Bài học Kinh nghiệm", "Trang 25"),
        ("PHẦN 6: CHỮ KÝ XÁC NHẬN VÀ BẢO MẬT DỮ LIỆU DOANH NGHIỆP", "Trang 26")
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        if title.startswith("PHẦN"):
            r1.bold = True
            r1.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_page_break()

    # -------------------------------------------------------------
    # PHẦN 1
    # -------------------------------------------------------------
    add_h1("PHẦN 1: CHƯƠNG TRÌNH THỰC TẬP VÀ THÔNG TIN DOANH NGHIỆP")
    
    add_h2("1.1. Giới thiệu Doanh nghiệp tiếp nhận (Amazon Web Services)")
    add_p("Amazon Web Services (AWS) là nền tảng điện toán đám mây hàng đầu thế giới, cung cấp hơn 200 dịch vụ toàn diện từ các trung tâm dữ liệu trên toàn cầu. Tại Việt Nam, công ty tiếp nhận sinh viên thực tập thuộc AWS Viet Nam Company Limited, phối hợp cùng cộng đồng AWS Study Group thông qua chương trình thực tập chuyên sâu First Cloud Journey (FCJ).")
    add_p("Chương trình FCJ AI/ML Track tập trung đào tạo và hướng dẫn sinh viên áp dụng các dịch vụ Cloud AI/ML cao cấp như Amazon SageMaker, AWS Lambda, Amazon API Gateway, Amazon S3, và Amazon CloudWatch để giải quyết các bài toán thực tế của doanh nghiệp.")

    add_h2("1.2. Chương trình Thực tập ngoài trường đã được Khoa duyệt (Form D2 / Form D3)")
    add_p("Chương trình thực tập của sinh viên đã được Khoa Khoa học và Kỹ thuật Máy tính - Trường Đại học Bách Khoa TP.HCM phê duyệt chính thức theo Form D2 (và danh sách xét tuyển Form D3). Các nội dung công việc cam kết thực hiện bao gồm:")
    
    bullets = [
        "Nghiên cứu và xây dựng hệ thống dự báo doanh số bán lẻ dựa trên tập dữ liệu chuẩn Rossmann Store Sales.",
        "Thiết kế quy trình tiền xử lý dữ liệu và trích xuất 22 đặc trưng chuỗi thời gian (Calendar, Rolling Means, Lag Features).",
        "Huấn luyện, tối ưu và đánh giá các kiến trúc mô hình Machine Learning / Deep Learning (XGBoost Regressor vs PyTorch LSTM).",
        "Triển khai hệ thống Experiment Tracking với Amazon SageMaker Experiments và quy trình tự động hóa MLOps với SageMaker Pipelines.",
        "Xây dựng hạ tầng Serverless REST API (AWS Lambda + API Gateway) và giao diện trực quan Live Web UI Dashboard phục vụ dự báo thời gian thực.",
        "Viết các bài báo khoa học kỹ thuật (Technical Blogs) chia sẻ cho cộng đồng và xây dựng tài liệu thực hành Workshop hoàn chỉnh."
    ]
    for b in bullets:
        add_p(f"• {b}", space_after=3)

    # -------------------------------------------------------------
    # PHẦN 2
    # -------------------------------------------------------------
    add_h1("PHẦN 2: NỘI DUNG VÀ NHẬT KÝ HOẠT ĐỘNG THỰC TẬP (WORKLOG)")

    add_h2("2.1. Nhật ký hoạt động chi tiết từng tuần (Tuần 1 → Tuần 8)")
    add_p("Thời gian thực tập chính thức: 06/06/2026 → 15/08/2026 (Tổng cộng 8 tuần thực tế). Dưới đây là nhật ký công việc chi tiết được ghi nhận theo từng tuần:")

    worklog_data = [
        ("Tuần 1 (06/06 - 12/06/2026)", "Onboarding & Thu thập Dữ liệu", "Tham gia buổi định hướng FCJ Program; Khởi tạo môi trường AWS; Khám phá bộ dữ liệu Rossmann Store Sales (1,017,209 dòng × 9 cột).", "Đã khởi tạo S3 Bucket `aws-internship-hkq-2026` và tải thành công dữ liệu thô lên S3."),
        ("Tuần 2 (13/06 - 19/06/2026)", "Tiền xử lý & Feature Engineering", "Xử lý dữ liệu khuyết, lọc cửa hàng đóng cửa (Open=0); Trích xuất 22 đặc trưng chuỗi thời gian (Rolling Means 7/14/30, Lag 7/14/30, Calendar features).", "Hoàn thành script `preprocessing.py`; Phân chia tập Train/Val/Test theo đúng trình tự thời gian."),
        ("Tuần 3 (20/06 - 26/06/2026)", "Huấn luyện Mô hình XGBoost Baseline", "Xây dựng kịch bản huấn luyện XGBoost Regressor (v1.7.6); Tối ưu siêu tham số bằng Optuna; Phân tích lỗi theo RMSE và MAPE.", "Mô hình XGBoost đạt kết quả vượt trội: Test RMSE = 925.28, Test MAPE = 9.92%."),
        ("Tuần 4 (27/06 - 03/07/2026)", "Thử nghiệm PyTorch LSTM & So sánh", "Thiết kế kiến trúc Deep Learning PyTorch LSTM (2 lớp LSTM + Linear output); Tiến hành huấn luyện trên GPU; So sánh trực diện với XGBoost.", "XGBoost áp đảo LSTM (MAPE 9.92% vs 32.79%); Quyết định lựa chọn XGBoost làm Production Model."),
        ("Tuần 5 (04/07 - 10/07/2026)", "SageMaker Experiments & SHAP Analysis", "Tích hợp thư viện `boto3` để ghi nhận siêu tham số và chỉ số đánh giá lên Amazon SageMaker Experiments; Trực quan hóa SHAP Values.", "Hoàn thành bài viết Blog 2 về SageMaker Experiments (Tác giả: Nguyễn Ngọc Sáng); Xác định Top 5 đặc trưng quan trọng."),
        ("Tuần 6 (11/07 - 17/07/2026)", "SageMaker Pipelines & Serving API", "Đóng gói quy trình MLOps tự động hóa với SageMaker Pipelines; Triển khai SageMaker Endpoint (`ml.t2.medium`); Triển khai AWS Lambda + API Gateway.", "Hoàn thành bài viết Blog 3 về MLOps Pipelines (Tác giả: Văn Thái Quân); Xây dựng xong Serverless REST API."),
        ("Tuần 7 (18/07 - 24/07/2026)", "Live Web Dashboard & Quality Gate", "Thiết kế giao diện Live Web UI Dashboard (Dark Mode / Glassmorphism); Kiểm thử giả lập kịch bản What-If; Đánh giá độ trễ API.", "API kiểm thử trên dữ liệu thực tế đạt độ chính xác ấn tượng (sai số 4.58%, độ trễ ~1.1s)."),
        ("Tuần 8 (25/07 - 15/08/2026)", "Tài liệu Workshop & Hoàn thiện Báo cáo", "Tổng hợp toàn bộ tài liệu hướng dẫn thực hành Workshop `Workshop_AWS_ML_Forecasting.md`; Đánh giá bản thân và hoàn thiện hồ sơ thực tập.", "Hoàn thành 100% các hạng mục công việc; Deploy toàn bộ sản phẩm lên GitHub và GitHub Pages.")
    ]

    wl_table = doc.add_table(rows=len(worklog_data) + 1, cols=4)
    wl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    wl_table.autofit = False
    
    headers = ["Tuần / Thời gian", "Tên Giai đoạn", "Mô tả Công việc Thực hiện", "Kết quả Đạt được"]
    hdr_widths = [Inches(1.5), Inches(1.3), Inches(2.3), Inches(1.8)]

    # Style Header Row
    hdr_row = wl_table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = hdr_widths[idx]
        set_cell_background(cell, "003366") # Navy
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Style Data Rows
    for r_idx, (w_time, w_stage, w_desc, w_res) in enumerate(worklog_data):
        row = wl_table.rows[r_idx + 1]
        bg_color = "F2F5F8" if r_idx % 2 == 1 else "FFFFFF"
        
        vals = [w_time, w_stage, w_desc, w_res]
        for c_idx, val in enumerate(vals):
            cell = row.cells[c_idx]
            cell.width = hdr_widths[c_idx]
            set_cell_background(cell, bg_color)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if c_idx == 0:
                run.bold = True

    add_p("", space_after=12)

    # -------------------------------------------------------------
    # PHẦN 3
    # -------------------------------------------------------------
    add_h1("PHẦN 3: BÁO CÁO KỸ THUẬT VÀ KIẾN TRÚC GIẢI PHÁP")

    add_h2("3.1. Tổng quan Kiến trúc MLOps 4 Tầng")
    add_p("Hệ thống được thiết kế theo kiến trúc MLOps 4 tầng độc lập, đảm bảo khả năng mở rộng, tính tin cậy và tuân thủ các tiêu chuẩn kỹ thuật doanh nghiệp:")

    add_p("1. Tầng Dữ liệu (Data Lake Layer): Amazon S3 đóng vai trò kho lưu trữ tập trung dữ liệu thô (`raw/`), dữ liệu sau tiền xử lý (`processed/`) và sản phẩm mô hình (`models/`).")
    add_p("2. Tầng Machine Learning (Training Layer): Tiến trình tiền xử lý 22 đặc trưng chuỗi thời gian, huấn luyện mô hình XGBoost Regressor và ghi vết thí nghiệm tự động lên Amazon SageMaker Experiments.")
    add_p("3. Tầng Phục vụ (Serving Layer): Mô hình được triển khai lên SageMaker Endpoint (`ml.t2.medium`), kết hợp với AWS Lambda wrapper và Amazon API Gateway để cung cấp chuẩn giao tiếp REST API công khai.")
    add_p("4. Tầng Giám sát (Monitoring Layer): Tích hợp thuật toán kiểm tra trôi dữ liệu (Z-Score Data Drift Detection) và hiển thị trực quan trên Amazon CloudWatch Dashboard.")

    add_h2("3.2. Tiền xử lý Dữ liệu và Feature Engineering (22 đặc trưng)")
    add_p("Bộ dữ liệu Rossmann Store Sales bao gồm 1,017,209 dòng giao dịch từ 1,115 cửa hàng. Tiến trình tiền xử lý loại bỏ 172,817 dòng khi cửa hàng đóng cửa (`Open = 0`) và trích xuất 22 đặc trưng kỹ thuật:")

    feat_items = [
        "Đặc trưng Lịch (Calendar Features): Year, Month, Day, DayOfWeek, WeekOfYear, IsWeekend, IsDecember.",
        "Trung bình Trượt (Rolling Means): rolling_mean_7, rolling_mean_14, rolling_mean_30.",
        "Độ trễ Thời gian (Lag Features): sales_lag_7, sales_lag_14, sales_lag_30.",
        "Đặc trưng Cửa hàng & Khuyến mại: StoreType, Assortment, CompetitionDistance, Promo, Promo2, StateHoliday, SchoolHoliday."
    ]
    for fi in feat_items:
        add_p(f"• {fi}", space_after=2)

    add_h2("3.3. Huấn luyện và So sánh Mô hình (XGBoost vs PyTorch LSTM)")
    add_p("Nhóm thực tập đã triển khai huấn luyện thực nghiệm và so sánh đối chứng giữa 2 thuật toán đại diện cho Học máy truyền thống (XGBoost) và Học sâu (PyTorch LSTM):")

    comp_table = doc.add_table(rows=3, cols=5)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_table.autofit = False
    
    comp_headers = ["Thuật toán Mô hình", "Test RMSE", "Test MAPE", "Thời gian Huấn luyện", "Đánh giá & Quyết định"]
    comp_widths = [Inches(1.8), Inches(1.1), Inches(1.1), Inches(1.4), Inches(1.5)]

    for idx, text in enumerate(comp_headers):
        cell = comp_table.rows[0].cells[idx]
        cell.width = comp_widths[idx]
        set_cell_background(cell, "003366")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)

    comp_rows = [
        ("XGBoost Regressor (v1.7.6)", "925.28", "9.92%", "~45 giây", "✅ Lựa chọn Production"),
        ("PyTorch LSTM (2-layer)", "3,044.43", "32.79%", "~8 phút", "❌ Thử nghiệm không đạt")
    ]
    for r_i, r_data in enumerate(comp_rows):
        row = comp_table.rows[r_i + 1]
        for c_i, val in enumerate(r_data):
            cell = row.cells[c_i]
            cell.width = comp_widths[c_i]
            set_cell_background(cell, "F2F5F8" if r_i == 1 else "FFFFFF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if c_i == 0 or c_i == 4:
                run.bold = True

    add_p("", space_after=6)
    add_p("Kết luận chuyên môn: Đối với dữ liệu chuỗi thời gian dạng bảng (Tabular Time Series), thuật toán XGBoost kết hợp với đặc trưng Rolling và Lag mang lại độ chính xác vượt trội (MAPE 9.92% so với 32.79% của LSTM) và tối ưu vượt bậc về chi phí tính toán.")

    add_h2("3.4. Quản lý Thí nghiệm với Amazon SageMaker Experiments (Blog 2 - Nguyễn Ngọc Sáng)")
    add_p("Bài viết kỹ thuật Blog 2 do thành viên Nguyễn Ngọc Sáng thực hiện trình bày giải pháp ghi vết thí nghiệm Machine Learning bằng Amazon SageMaker Experiments. Giải pháp cho phép nhóm nghiên cứu lưu trữ toàn bộ lịch sử siêu tham số, chỉ số RMSE/MAPE và trực quan hóa đường cong học tập (Learning Curves) trực tiếp qua giao diện AWS Console thông qua kết nối API `boto3` từ môi trường local.")

    add_h2("3.5. Tự động hóa Pipeline với Amazon SageMaker Pipelines (Blog 3 - Văn Thái Quân)")
    add_p("Bài viết kỹ thuật Blog 3 do thành viên Văn Thái Quân thực hiện trình bày kiến trúc tự động hóa quy trình MLOps bằng SageMaker Pipelines. Mã nguồn tiền xử lý và cấu hình huấn luyện được đóng gói thành các tập tin `sourcedir.tar.gz` lưu trữ trên S3, giúp quy trình huấn luyện tự động cấp phát tài nguyên, thực thi và lưu trữ mô hình một cách độc lập và an toàn.")

    add_h2("3.6. Triển khai Serverless REST API và Live Web UI Dashboard")
    add_p("Để cung cấp khả năng dự báo thời gian thực cho ứng dụng phía người dùng cuối, hệ thống sử dụng AWS Lambda đóng vai trò wrapper nhận yêu cầu REST từ Amazon API Gateway và gọi API `sagemaker-runtime.invoke_endpoint`. Mô hình được kiểm thử trên dữ liệu thực tế (Store 1, ngày 2015-06-15) đạt độ chính xác ấn tượng với mức sai lệch chỉ 4.58% (Doanh số thực tế: 5,518 vs Doanh số dự báo: 5,770.64).")
    add_p("Nhóm cũng hoàn thiện giao diện Live Web UI Dashboard (Dark Mode / Glassmorphism) chạy tại cổng 8000 hỗ trợ trực quan hóa biểu đồ xu hướng 14 ngày và mô phỏng kịch bản What-If.")

    # -------------------------------------------------------------
    # PHẦN 4
    # -------------------------------------------------------------
    add_h1("PHẦN 4: HƯỚNG DẪN THỰC HÀNH WORKSHOP PIPELINE")
    add_p("Toàn bộ quy trình thực hành xây dựng Pipeline Dự báo Doanh số trên AWS đã được đóng gói chi tiết trong tập tin Workshop_AWS_ML_Forecasting.md. Các bước thực hiện chính bao gồm:")

    ws_steps = [
        "Bước 1: Khởi tạo IAM Role với các quyền truy cập tối thiểu cho S3, SageMaker, Lambda, API Gateway và CloudWatch.",
        "Bước 2: Cài đặt môi trường ảo Python local (`venv`) và cấu hình tập tin `config.py` kết nối với S3 Bucket.",
        "Bước 3: Thực thi script `preprocessing.py` để làm sạch dữ liệu, tạo 22 đặc trưng và phân chia tập dữ liệu theo mốc thời gian.",
        "Bước 4: Chạy kịch bản `train_xgboost.py` huấn luyện mô hình XGBoost và phân tích độ quan trọng đặc trưng với `shap_analysis.py`.",
        "Bước 5: Khởi tạo SageMaker Endpoint (`deploy_endpoint.py`), triển khai AWS Lambda (`deploy_lambda.py`) và khởi chạy Live UI Dashboard (`demo_ui/server.py`).",
        "Bước 6: Thực thi kịch bản dọn dẹp tài nguyên tự động `cleanup.py` để xóa SageMaker Endpoint nhằm tránh phát sinh chi phí phát sinh."
    ]
    for ws in ws_steps:
        add_p(f"• {ws}", space_after=3)

    # -------------------------------------------------------------
    # PHẦN 5: TỰ ĐÁNH GIÁ VÀ PHẢN HỒI THỰC TẬP (8 TUẦN)
    # -------------------------------------------------------------
    add_h1("PHẦN 5: TỰ ĐÁNH GIÁ VÀ PHẢN HỒI THỰC TẬP (8 TUẦN)")

    add_h2("5.1. Bảng Tiêu chí Đánh giá Cá nhân (Evaluation Criteria Table)")
    add_p("Dưới đây là bảng đánh giá 8 tiêu chí chuyên môn dựa trên các kết quả thực tế đạt được trong 8 tuần thực tập:")

    eval_table = doc.add_table(rows=9, cols=4)
    eval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    eval_table.autofit = False
    
    e_headers = ["#", "Tiêu chí Đánh giá", "Đánh giá", "Mô tả Chi tiết & Minh chứng Thực tế"]
    e_widths = [Inches(0.4), Inches(2.2), Inches(1.1), Inches(3.2)]

    for idx, text in enumerate(e_headers):
        cell = eval_table.rows[0].cells[idx]
        cell.width = e_widths[idx]
        set_cell_background(cell, "003366")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)

    eval_data = [
        ("1", "Kiến thức AWS Kỹ thuật", "Tốt", "Thành thạo S3, IAM, SageMaker Endpoints/Experiments/Pipelines, Lambda, API Gateway, CloudWatch."),
        ("2", "Kỹ năng Machine Learning", "Tốt", "Trích xuất 22 đặc trưng chuỗi thời gian; train XGBoost đạt Test RMSE 925.28 và MAPE 9.92%."),
        ("3", "Giải quyết Vấn đề & Debug", "Tốt", "Xử lý triệt để 3 sự cố: Nâng Service Quotas, ấn định version SDK `sagemaker==2.257.5`, sửa tràn số log `np.expm1()`."),
        ("4", "Chất lượng Code & Kiến trúc", "Khá", "Mã nguồn tổ chức mô-đun hóa sạch sẽ; đóng gói tự động `sourcedir.tar.gz`."),
        ("5", "Làm việc Nhóm & Hợp tác", "Tốt", "Phân công nhịp nhàng 3 vai trò (Data/ML, Backend, Infra); giao tiếp minh bạch và họp trao đổi định kỳ."),
        ("6", "Quản lý Thời gian", "Khá", "Hoàn thành 100% mục tiêu của 8 tuần thực tập; tiến độ các mốc chính được kiểm soát tốt."),
        ("7", "Tài liệu Kỹ thuật", "Tốt", "Đăng 3 bài blog kỹ thuật học thuật; viết file tài liệu thực hành Workshop `Workshop_AWS_ML_Forecasting.md`."),
        ("8", "Chủ động & Sáng kiến", "Tốt", "Chủ động kiểm tra quota sớm; đề xuất kiến trúc Serverless REST API với Lambda + API Gateway.")
    ]

    for r_i, r_data in enumerate(eval_data):
        row = eval_table.rows[r_i + 1]
        bg_c = "F2F5F8" if r_i % 2 == 1 else "FFFFFF"
        for c_i, val in enumerate(r_data):
            cell = row.cells[c_i]
            cell.width = e_widths[c_i]
            set_cell_background(cell, bg_c)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if c_i == 0 or c_i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True

    add_p("", space_after=10)

    add_h2("5.2. Đánh giá Tổng quan & Phản hồi Thực tập")
    add_h3("1. Môi trường Làm việc (Working Environment)")
    add_p("Môi trường làm việc rất thân thiện và cởi mở. Các thành viên FCAJ luôn sẵn sàng hỗ trợ mỗi khi tôi gặp khó khăn, ngay cả ngoài giờ làm việc. Không gian làm việc ngăn nắp, thoải mái, giúp tôi tập trung tốt hơn.")

    add_h3("2. Sự Hỗ trợ từ Mentor / Ban Quản lý (Support from Mentor / Team Admin)")
    add_p("Mentor hướng dẫn rất chi tiết, giải thích rõ ràng khi tôi chưa hiểu và luôn khuyến khích tôi đặt câu hỏi. Đội ngũ Admin hỗ trợ hiệu quả các thủ tục hành chính, cung cấp đầy đủ tài liệu cần thiết và tạo điều kiện thuận lợi nhất để tôi làm việc.")

    add_h3("3. Mức độ Phù hợp với Chuyên ngành (Relevance of Work to Academic Major)")
    add_p("Các công việc được giao bám sát kiến thức tôi đã học tại trường đại học, đồng thời giới thiệu thêm nhiều mảng công nghệ mới giúp tôi vừa củng cố nền tảng vừa tích lũy được nhiều kỹ năng thực tế có giá trị cao.")

    add_h3("4. Cơ hội Học hỏi & Phát triển Kỹ năng (Learning & Skill Development Opportunities)")
    add_p("Trong suốt 8 tuần thực tập, tôi đã học hỏi được nhiều kỹ năng mới như sử dụng công cụ quản lý dự án, kỹ năng làm việc nhóm và giao tiếp chuyên nghiệp trong môi trường doanh nghiệp.")

    add_h3("5. Văn hóa Công ty & Tinh thần Đồng đội (Company Culture & Team Spirit)")
    add_p("Văn hóa công ty rất tích cực: mọi người tôn trọng lẫn nhau, làm việc nghiêm túc nhưng vẫn giữ không khí vui vẻ, hỗ trợ lẫn nhau không phân biệt vị trí.")

    add_h3("6. Chính sách & Quyền lợi Thực tập (Internship Policies / Benefits)")
    add_p("Công ty cung cấp phụ cấp thực tập và hỗ trợ thời gian làm việc linh hoạt khi cần thiết. Ngoài ra, cơ hội tham gia các buổi đào tạo nội bộ chuyên sâu là một điểm cộng rất lớn của chương trình.")

    add_h2("5.3. Đề xuất & Bài học Kinh nghiệm (Recommendations)")
    add_p("• Kiểm tra Service Quotas từ Ngày 1: Luôn kiểm tra quota AWS trước khi viết bất kỳ mã nguồn huấn luyện nào.")
    add_p("• Sử dụng API boto3 trực tiếp: Tăng tính ổn định cho các luồng công việc phức tạp trên SageMaker.")
    add_p("• Xây dựng Knowledge Base nội bộ: Đề xuất công ty xây dựng kho kiến thức lưu trữ cách xử lý các sự cố tài nguyên đám mây thường gặp để giúp thực tập sinh khóa sau hòa nhập nhanh chóng hơn.")

    # -------------------------------------------------------------
    # PHẦN 6: CHỮ KÝ VÀ XÁC NHẬN
    # -------------------------------------------------------------
    add_h1("PHẦN 6: CHỮ KÝ XÁC NHẬN VÀ BẢO MẬT DỮ LIỆU DOANH NGHIỆP")
    add_p("Báo cáo này đã được đại diện Doanh nghiệp tiếp nhận thực tập kiểm tra, xác nhận không vi phạm các quy định về bảo mật thông tin (NDA) và chấp thuận cho sinh viên nộp về Khoa Khoa học và Kỹ thuật Máy tính - Trường Đại học Bách Khoa TP.HCM.")

    add_p("", space_after=20)

    # Signature Table
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    sig_widths = [Inches(3.4), Inches(3.4)]
    
    # Header Cell Left
    c_left_hdr = sig_table.rows[0].cells[0]
    c_left_hdr.width = sig_widths[0]
    p_l = c_left_hdr.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l.paragraph_format.space_after = Pt(2)
    r = p_l.add_run("XÁC NHẬN CỦA ĐẠI DIỆN DOANH NGHIỆP\n(Ký tên, ghi rõ họ tên và đóng dấu)")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    # Header Cell Right
    c_right_hdr = sig_table.rows[0].cells[1]
    c_right_hdr.width = sig_widths[1]
    p_r = c_right_hdr.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r.paragraph_format.space_after = Pt(2)
    r = p_r.add_run("SINH VIÊN THỰC HIỆN BÁO CÁO\n(Ký và ghi rõ họ tên - Chữ ký màu xanh)")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    # Content Row (Space for signature)
    c_left_body = sig_table.rows[1].cells[0]
    c_left_body.width = sig_widths[0]
    p_lb = c_left_body.paragraphs[0]
    p_lb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lb.paragraph_format.space_before = Pt(60)
    r_lb = p_lb.add_run("CÁN BỘ HƯỚNG DẪN DOANH NGHIỆP\nNguyễn Ngọc Sáng / Nguyễn Đức Hoàng")
    r_lb.bold = True
    r_lb.font.name = 'Times New Roman'
    r_lb.font.size = Pt(12)

    c_right_body = sig_table.rows[1].cells[1]
    c_right_body.width = sig_widths[1]
    p_rb = c_right_body.paragraphs[0]
    p_rb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rb.paragraph_format.space_before = Pt(60)
    r_rb = p_rb.add_run("SINH VIÊN THỰC TẬP\n\nHuỳnh Kim Quý\nMSSV: 2312918")
    r_rb.bold = True
    r_rb.font.name = 'Times New Roman'
    r_rb.font.size = Pt(12)

    output_path = r"e:\AWS - TTNT\BaoCao_ThucTap_NgoaiTruong_AWS_SageMaker_MLOps.docx"
    doc.save(output_path)
    print(f"Successfully generated updated Word report at: {output_path}")

if __name__ == "__main__":
    create_report()
